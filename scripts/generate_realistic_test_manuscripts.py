from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent.parent


@dataclass(frozen=True)
class AuthorFixture:
    name: str
    email: str
    affiliation: str
    city: str
    country_or_region: str
    is_corresponding: bool = False


@dataclass(frozen=True)
class ManuscriptFixture:
    slug: str
    title: str
    journal_hint: str
    abstract: str
    keywords: list[str]
    authors: list[AuthorFixture]
    sections: list[tuple[str, str]]


FIXTURES: list[ManuscriptFixture] = [
    ManuscriptFixture(
        slug="academic_paper_01",
        title="Digital Twin Governance for Cross-Border Cold-Chain Logistics",
        journal_hint="Computer Science",
        abstract=(
            "Digital twin deployments in cross-border cold-chain logistics remain fragmented across customs, "
            "warehouse, and carrier systems. This study proposes a governance architecture that coordinates "
            "sensor provenance, alert escalation, and traceability obligations across international handoff points. "
            "Using three simulated trade corridors, we show that the proposed policy layer reduces unverified temperature "
            "events while preserving operational response speed."
        ),
        keywords=["digital twin", "cold chain", "governance", "traceability", "cross-border logistics"],
        authors=[
            AuthorFixture("Alice Chen", "alice.chen@ccnu.edu.cn", "Central China Normal University", "Wuhan", "China", True),
            AuthorFixture("Bob Li", "bob.li@whu.edu.cn", "Wuhan University", "Wuhan", "China"),
            AuthorFixture("Maya Thompson", "maya.thompson@qmul.ac.uk", "Queen Mary University of London", "London", "United Kingdom"),
        ],
        sections=[
            ("1. Introduction", "Cross-border cold-chain logistics involve multiple operational systems that frequently disagree on source-of-truth temperature events and responsibility boundaries."),
            ("2. Method", "We modelled a policy-first digital twin architecture and evaluated it against baseline workflows across three representative logistics corridors."),
            ("3. Results", "The governance layer reduced unverified alerts, shortened exception handoff time, and improved audit completeness for downstream quality review."),
            ("4. Conclusion", "A shared policy layer can make digital twins more trustworthy in operationally fragmented cold-chain environments."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_02",
        title="Adaptive Prompt Routing for Resource-Constrained Scholarly Assistants",
        journal_hint="Computer Science",
        abstract=(
            "Scholarly assistants increasingly rely on multiple model backends with heterogeneous latency and cost profiles. "
            "We introduce an adaptive prompt routing strategy that classifies editorial subtasks into extraction, verification, "
            "and drafting bands, and dispatches them to different models under a budget-aware controller. "
            "Across twelve journal-office scenarios, the controller preserves factual consistency while cutting median response cost."
        ),
        keywords=["prompt routing", "editorial AI", "latency control", "budget-aware inference"],
        authors=[
            AuthorFixture("Jeffrey Lu", "jeffrey.lu@scholarflow.ai", "ScholarFlow Research Lab", "Shanghai", "China", True),
            AuthorFixture("Sara Nabil", "s.nabil@manchester.ac.uk", "The University of Manchester", "Manchester", "United Kingdom", True),
            AuthorFixture("Tianyu Zhao", "tianyu.zhao@zju.edu.cn", "Zhejiang University", "Hangzhou", "China"),
        ],
        sections=[
            ("1. Introduction", "Editorial assistants must balance throughput, factuality, and compute budgets across a wide variety of manuscript operations."),
            ("2. Routing Policy", "We define a three-band classifier and combine it with runtime quotas, confidence thresholds, and mandatory verification stages."),
            ("3. Evaluation", "The adaptive strategy reduced unnecessary premium-model calls without harming extraction accuracy in the benchmark tasks."),
            ("4. Discussion", "Budget-aware routing is viable only when failure modes are explicit and fallback policies are deterministic."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_03",
        title="Multisite Evidence on Faculty Adoption of AI Writing Feedback Tools",
        journal_hint="Education",
        abstract=(
            "Universities are rapidly deploying AI writing feedback tools, yet faculty adoption patterns remain uneven. "
            "This multisite study compares policy framing, assessment redesign, and trust-building practices across five institutions. "
            "We find that adoption depends less on headline accuracy claims and more on local governance, exemplar review workflows, and student appeal pathways."
        ),
        keywords=["AI feedback", "higher education", "faculty adoption", "assessment redesign"],
        authors=[
            AuthorFixture("Elena García", "e.garcia@ub.edu", "Universitat de Barcelona", "Barcelona", "Spain", True),
            AuthorFixture("Rui Wang", "rui.wang@pku.edu.cn", "Peking University", "Beijing", "China"),
            AuthorFixture("Jin Chang", "jin.chang@yonsei.ac.kr", "Yonsei University", "Seoul", "South Korea"),
        ],
        sections=[
            ("1. Background", "Institutional roll-out of AI writing support tools often outpaces shared norms for assessment and appeals."),
            ("2. Data Collection", "We conducted interviews, policy audits, and workshop observations across five campuses over two semesters."),
            ("3. Findings", "Faculty trust increased when reviewers could inspect system rationale and retain authority over contested outputs."),
            ("4. Implications", "Adoption programmes should prioritise policy clarity and exemplar-based onboarding rather than marketing claims."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_04",
        title="Explainable Risk Signals for Mid-Sized Hospital Procurement Networks",
        journal_hint="Healthcare Management",
        abstract=(
            "Procurement teams in mid-sized hospitals face volatile shortages without access to advanced planning systems. "
            "We design an explainable risk-signal pipeline that combines supplier concentration indicators, contract age, and logistics fragility into a weekly alert surface. "
            "The approach improves planner confidence because each alert is paired with human-readable drivers and suggested mitigation actions."
        ),
        keywords=["procurement risk", "hospital operations", "explainable analytics", "supply resilience"],
        authors=[
            AuthorFixture("Priya Raman", "priya.raman@nus.edu.sg", "National University of Singapore", "Singapore", "Singapore", True),
            AuthorFixture("Dong Li", "dong.li@fudan.edu.cn", "Fudan University", "Shanghai", "China"),
            AuthorFixture("Wameed Alghazali", "w.alghazali@uaeu.ac.ae", "United Arab Emirates University", "Al Ain", "United Arab Emirates"),
        ],
        sections=[
            ("1. Problem Statement", "Hospital procurement teams require decision support that remains intelligible to clinical and administrative stakeholders."),
            ("2. Signal Design", "The pipeline composes supplier concentration, contract age, and route fragility metrics into transparent weekly alerts."),
            ("3. Operational Study", "Pilot teams acted on alerts more frequently when each recommendation included plain-language justification."),
            ("4. Conclusion", "Explainability is not ornamental in hospital procurement; it is a prerequisite for operational adoption."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_05",
        title="Revisiting Community Microgrids with Coordinated Storage and Demand Response",
        journal_hint="Energy Systems",
        abstract=(
            "Community microgrids promise resilience, yet many deployments underperform because storage and demand response are planned independently. "
            "We evaluate a coordinated dispatch policy for neighbourhood-scale microgrids under mixed outage and peak-pricing scenarios. "
            "The combined strategy improves resilience metrics and lowers emergency curtailment compared with storage-only baselines."
        ),
        keywords=["microgrid", "demand response", "energy storage", "resilience"],
        authors=[
            AuthorFixture("Mariola Wasil", "mariola.wasil@pw.edu.pl", "Warsaw University of Technology", "Warsaw", "Poland", True),
            AuthorFixture("Luke Griffiths", "luke.griffiths@ngi.no", "Norwegian Geotechnical Institute", "Oslo", "Norway"),
        ],
        sections=[
            ("1. Introduction", "Storage assets alone do not guarantee microgrid resilience when residential demand is poorly coordinated."),
            ("2. Scenario Design", "We model outage clusters, dynamic tariffs, and appliance flexibility in a neighbourhood-scale simulation."),
            ("3. Results", "Coordinated storage and demand response reduced forced curtailment and improved recovery time after outages."),
            ("4. Conclusion", "Microgrid planning should treat flexible demand as a first-class resilience resource."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_06",
        title="Public Reasoning Quality in Multilingual Municipal Consultation Portals",
        journal_hint="Public Policy",
        abstract=(
            "Municipal consultation portals increasingly rely on multilingual participation, but moderation and synthesis tools often privilege a dominant language. "
            "We propose a multilingual reasoning-quality rubric and apply it to three consultation platforms in Asia and Europe. "
            "The study shows that interface translation alone is insufficient; quality improves when argument summaries preserve stance, evidence, and local context."
        ),
        keywords=["public consultation", "multilingual systems", "reasoning quality", "municipal governance"],
        authors=[
            AuthorFixture("Fei Liu", "fei.liu@hkbu.edu.hk", "Hong Kong Baptist University", "Hong Kong", "Hong Kong SAR", True),
            AuthorFixture("Fernando Marques", "fsmarques@fc.ul.pt", "Universidade de Lisboa", "Lisbon", "Portugal"),
            AuthorFixture("Matthew Coop", "m.coop@ucl.ac.uk", "University College London", "London", "United Kingdom"),
        ],
        sections=[
            ("1. Context", "Multilingual civic platforms often offer translation but fail to preserve evidentiary nuance in public reasoning."),
            ("2. Rubric Construction", "Our rubric evaluates stance clarity, evidence retention, contextual fidelity, and moderation traceability."),
            ("3. Comparative Analysis", "The highest-scoring platform paired translation with editor review of argument summaries."),
            ("4. Implications", "Municipal consultation design should measure reasoning preservation rather than surface language coverage alone."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_07",
        title="Trust Repair After Data Pipeline Incidents in Research Software Teams",
        journal_hint="Software Engineering",
        abstract=(
            "Research software teams often frame data pipeline incidents as technical mishaps, yet trust erosion after a breach or silent data corruption has organisational roots. "
            "This paper analyses trust-repair practices across eight software teams supporting academic labs. "
            "The strongest repair outcomes emerged when incident reviews explicitly linked code fixes, communication repair, and reproducibility checks."
        ),
        keywords=["trust repair", "research software", "incident response", "reproducibility"],
        authors=[
            AuthorFixture("Kaiping Zheng", "kaiping.zheng@whut.edu.cn", "Wuhan University of Technology", "Wuhan", "China", True),
            AuthorFixture("Kazimierz Józefiak", "k.jozefiak@il.pw.edu.pl", "Warsaw University of Technology", "Warsaw", "Poland"),
            AuthorFixture("Srevan Muguda", "srevan.muguda-viswanath@durham.ac.uk", "Durham University", "Durham", "United Kingdom"),
        ],
        sections=[
            ("1. Introduction", "Pipeline incidents harm confidence in both data products and the teams that maintain them."),
            ("2. Study Design", "We analysed post-incident artefacts, interview data, and repair strategies across eight research software teams."),
            ("3. Findings", "Trust repair improved when teams paired code remediation with explicit reproducibility rechecks and stakeholder communication."),
            ("4. Conclusion", "Incident recovery in research software is socio-technical work, not only a patching exercise."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_08",
        title="Low-Resource OCR Benchmarking for Archival Humanities Collections",
        journal_hint="Digital Humanities",
        abstract=(
            "Small archives often digitise multilingual holdings without the computational budget required for large OCR pipelines. "
            "We benchmark low-resource OCR configurations across four archival collections containing mixed print quality and orthography variation. "
            "The results highlight the trade-off between recogniser flexibility, post-correction effort, and cataloguing throughput."
        ),
        keywords=["OCR", "digital humanities", "archives", "benchmarking"],
        authors=[
            AuthorFixture("Arif Mohammad", "arif.mohammad@du.ac.bd", "University of Dhaka", "Dhaka", "Bangladesh", True),
            AuthorFixture("Tiantian Ma", "tiantian.ma@sjtu.edu.cn", "Shanghai Jiao Tong University", "Shanghai", "China"),
        ],
        sections=[
            ("1. Background", "Archival OCR in low-resource settings must balance recogniser quality against staffing and correction budgets."),
            ("2. Benchmark Setup", "We compared four configurations over mixed-language archival scans with uneven print quality."),
            ("3. Findings", "Lower-cost recognisers remained competitive when paired with targeted post-correction routines."),
            ("4. Conclusion", "Practical OCR strategies for archives should be optimised for sustainable correction workflows, not only raw character accuracy."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_09",
        title="Quantifying Editorial Delay Drivers in Fast-Growth Interdisciplinary Journals",
        journal_hint="Publishing Studies",
        abstract=(
            "Fast-growth journals often struggle to preserve predictable editorial turnaround as submission volume diversifies. "
            "We model delay drivers across intake, reviewer assignment, revision, and decision stages in an interdisciplinary publishing programme. "
            "The analysis identifies reviewer outreach quality, missing manuscript metadata, and repeated role handoffs as the dominant contributors to elapsed time."
        ),
        keywords=["editorial workflow", "turnaround time", "journal operations", "process analytics"],
        authors=[
            AuthorFixture("Louis Shawn", "louis.shawn@qq.com", "ScholarFlow Operations Lab", "Wuhan", "China", True),
            AuthorFixture("Azhar Abbas", "azhar.abbas@zalf.de", "Leibniz Centre for Agricultural Landscape Research", "Müncheberg", "Germany"),
            AuthorFixture("Jeffrey Lu", "lujeffrey45@gmail.com", "ScholarFlow Operations Lab", "Wuhan", "China"),
        ],
        sections=[
            ("1. Introduction", "Submission growth can amplify latent workflow bottlenecks across editorial intake and review coordination."),
            ("2. Data and Measures", "We assembled manuscript-stage timestamps and classified delay sources using queue analytics and manual review."),
            ("3. Results", "Reviewer outreach quality and missing manuscript metadata explained more delay than raw submission volume alone."),
            ("4. Conclusion", "Editorial operations improve fastest when intake completeness and reviewer invitation quality are treated as first-order levers."),
        ],
    ),
    ManuscriptFixture(
        slug="academic_paper_10",
        title="Scenario-Based Evaluation of Reviewer Recruitment Strategies for New Journals",
        journal_hint="Publishing Studies",
        abstract=(
            "New journals frequently overestimate the availability of willing reviewers, especially outside established editorial networks. "
            "This paper evaluates scenario-based reviewer recruitment strategies under varying novelty, turnaround, and policy constraints. "
            "We show that shortlist quality and invitation sequencing matter more than expanding the candidate pool indiscriminately."
        ),
        keywords=["reviewer recruitment", "journal launch", "editorial policy", "scenario planning"],
        authors=[
            AuthorFixture("ReviewerLouis", "1397951685@qq.com", "Independent Editorial Consultant", "Wuhan", "China", True),
            AuthorFixture("OwnerTest", "test@owner.com", "Owner Research Consulting", "Shenzhen", "China"),
            AuthorFixture("CSME", "new_editor@university.edu", "University Editorial Office", "Nanjing", "China"),
        ],
        sections=[
            ("1. Problem Framing", "Reviewer recruitment is often treated as a volume problem, even though sequencing and fit are more decisive."),
            ("2. Scenario Planning", "We compare conservative, balanced, and aggressive invitation strategies under varying deadline and quality assumptions."),
            ("3. Results", "The balanced strategy outperformed wider shotgun recruitment by preserving response quality and reducing cooldown collisions."),
            ("4. Practical Guidance", "Editorial teams should invest in shortlist quality and invitation orchestration before expanding raw candidate counts."),
        ],
    ),
]


def _author_superscript_map(authors: list[AuthorFixture]) -> tuple[list[int], dict[str, int]]:
    mapping: dict[str, int] = {}
    for author in authors:
        key = "|".join([author.affiliation, author.city, author.country_or_region])
        if key not in mapping:
            mapping[key] = len(mapping) + 1
    ordered = [mapping["|".join([a.affiliation, a.city, a.country_or_region])] for a in authors]
    return ordered, mapping


def build_docx(fixture: ManuscriptFixture, output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(fixture.title)
    run.bold = True
    run.font.size = Pt(18)

    affiliation_numbers, affiliation_mapping = _author_superscript_map(fixture.authors)

    authors_paragraph = doc.add_paragraph()
    authors_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, (author, aff_no) in enumerate(zip(fixture.authors, affiliation_numbers)):
        if index > 0:
            authors_paragraph.add_run(", ")
        name_run = authors_paragraph.add_run(author.name)
        name_run.font.size = Pt(11)
        sup_run = authors_paragraph.add_run(str(aff_no))
        sup_run.font.superscript = True
        sup_run.font.size = Pt(8)
        if author.is_corresponding:
            corr_run = authors_paragraph.add_run("*")
            corr_run.font.superscript = True
            corr_run.font.size = Pt(8)

    aff_keys = list(affiliation_mapping.items())
    ordered_affiliations = sorted(aff_keys, key=lambda item: item[1])
    for raw_key, number in ordered_affiliations:
        affiliation, city, country = raw_key.split("|")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        number_run = paragraph.add_run(str(number))
        number_run.font.superscript = True
        number_run.font.size = Pt(8)
        text_run = paragraph.add_run(f" {affiliation}, {city}, {country}")
        text_run.font.size = Pt(10)

    corresponding_authors = [author for author in fixture.authors if author.is_corresponding]
    if corresponding_authors:
        corr_paragraph = doc.add_paragraph()
        corr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        corr_run = corr_paragraph.add_run(
            "Correspondence: "
            + "; ".join(
                f"{author.name} ({author.email})" for author in corresponding_authors
            )
        )
        corr_run.italic = True
        corr_run.font.size = Pt(10)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(f"Target Journal: {fixture.journal_hint}")
    meta_run.bold = True
    meta_run.font.size = Pt(10)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h_run = heading.add_run("Abstract")
    h_run.bold = True
    h_run.font.size = Pt(12)

    abstract = doc.add_paragraph(fixture.abstract)
    abstract.style = doc.styles["Normal"]
    for run in abstract.runs:
        run.font.size = Pt(11)

    keywords = doc.add_paragraph()
    kw_label = keywords.add_run("Keywords: ")
    kw_label.bold = True
    kw_label.font.size = Pt(10)
    kw_values = keywords.add_run(", ".join(fixture.keywords))
    kw_values.font.size = Pt(10)

    for section_title, section_body in fixture.sections:
        section_heading = doc.add_paragraph()
        section_heading_run = section_heading.add_run(section_title)
        section_heading_run.bold = True
        section_heading_run.font.size = Pt(12)
        body_paragraph = doc.add_paragraph(section_body)
        for run in body_paragraph.runs:
            run.font.size = Pt(11)

    doc.save(output_path)


def build_pdf(fixture: ManuscriptFixture, output_path: Path) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCenter",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.black,
        spaceAfter=10,
    )
    center_style = ParagraphStyle(
        "CenterMeta",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontName="Helvetica",
        fontSize=10.5,
        leading=13,
        spaceAfter=4,
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=14,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )

    aff_numbers, aff_mapping = _author_superscript_map(fixture.authors)
    story = [Paragraph(fixture.title, title_style)]

    author_fragments: list[str] = []
    for author, aff_no in zip(fixture.authors, aff_numbers):
        suffix = f"<super>{aff_no}</super>"
        if author.is_corresponding:
            suffix += "<super>*</super>"
        author_fragments.append(f"{author.name}{suffix}")
    story.append(Paragraph(", ".join(author_fragments), center_style))

    ordered_affiliations = sorted(aff_mapping.items(), key=lambda item: item[1])
    for raw_key, number in ordered_affiliations:
        affiliation, city, country = raw_key.split("|")
        story.append(Paragraph(f"<super>{number}</super> {affiliation}, {city}, {country}", center_style))

    corresponding_authors = [author for author in fixture.authors if author.is_corresponding]
    if corresponding_authors:
        story.append(
            Paragraph(
                "Correspondence: "
                + "; ".join(f"{author.name} ({author.email})" for author in corresponding_authors),
                center_style,
            )
        )

    story.append(Spacer(1, 8))
    story.append(Paragraph("Abstract", heading_style))
    story.append(Paragraph(fixture.abstract, body_style))
    story.append(
        Paragraph(
            "<b>Keywords:</b> " + ", ".join(fixture.keywords),
            body_style,
        )
    )
    story.append(Paragraph(f"<b>Target Journal:</b> {fixture.journal_hint}", body_style))

    for title, body in fixture.sections:
        story.append(Paragraph(title, heading_style))
        story.append(Paragraph(body, body_style))

    doc.build(story)


def main() -> None:
    for fixture in FIXTURES:
        build_docx(fixture, ROOT / f"{fixture.slug}.docx")
        build_pdf(fixture, ROOT / f"{fixture.slug}.pdf")
        print(f"generated {fixture.slug}.docx/.pdf")


if __name__ == "__main__":
    main()
